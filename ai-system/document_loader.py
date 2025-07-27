import os, re, hashlib, logging, PyPDF2, ebooklib, chardet
from typing import List, Dict, Optional
from pathlib import Path
import pandas as pd
from docx import Document
from ebooklib import epub
from bs4 import BeautifulSoup
from sentence_transformers import SentenceTransformer

from openai import OpenAI
from pinecone import Pinecone, ServerlessSpec
from dotenv import load_dotenv

load_dotenv(".env")

class DocumentLoader:
    """Завантажувач документів в Pinecone векторну базу"""
    
    def __init__(self, pinecone_index_name: str, auto_create_index: bool = True, dimension: int = 1024):
        """
        Ініціалізація DocumentLoader
        
        Args:
            pinecone_index_name: Назва індексу Pinecone
            auto_create_index: Автоматично створювати індекс якщо не існує
            dimension: Розмірність векторів (768 для multilingual-e5-base)
        """
        self.openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        # Pinecone
        self.pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        self.index_name = pinecone_index_name
        self.dimension = dimension
        
        # Ініціалізуємо індекс (створюємо якщо не існує)
        self.index = self._init_pinecone_index(auto_create_index)
        
        # Автоматично визначаємо модель embedding на основі розмірності індексу
        self._detect_embedding_model()
        
        # Налаштування
        self.chunk_size = 1000  # Розмір чанків
        self.chunk_overlap = 200  # Перекриття між чанками
        
        # Підтримувані формати
        self.supported_formats = {
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.doc': self._load_docx,
            '.xlsx': self._load_excel,
            '.xls': self._load_excel,
            '.csv': self._load_csv,
            '.epub': self._load_epub,
            '.txt': self._load_text,
            '.md': self._load_text,
            '.py': self._load_text,
            '.json': self._load_text,
            '.jsonl': self._load_text
        }    
    def _init_pinecone_index(self, auto_create: bool = True):
        """Ініціалізація Pinecone індексу з автоматичним створенням"""
        try:
            # Перевіряємо чи існує індекс
            existing_indexes = [idx.name for idx in self.pc.list_indexes()]
            
            if self.index_name in existing_indexes:
                return self.pc.Index(self.index_name)
            
            # Якщо індекс не існує і auto_create=True
            if auto_create:                
                # Визначаємо специфікацію для безкоштовного плану
                spec = ServerlessSpec(
                    cloud='aws',  # або 'gcp'
                    region='us-east-1'  # регіон для безкоштовного плану
                )
                
                # Створюємо індекс
                self.pc.create_index(
                    name=self.index_name,
                    dimension=self.dimension,  # Використовуємо передану розмірність
                    metric='cosine',  # Метрика подібності
                    spec=spec
                )
                
                # Чекаємо поки індекс буде готовий
                import time
                max_wait_time = 60  # максимум 60 секунд
                wait_time = 0
                
                while wait_time < max_wait_time:
                    try:
                        index = self.pc.Index(self.index_name)
                        return index
                    except Exception as e:
                        time.sleep(5)
                        wait_time += 5
                
                # Якщо не вдалося дочекатися
                raise Exception(f"Timeout waiting for index {self.index_name}")
                
            else:
                raise Exception(f"Індекс '{self.index_name}' не існує і auto_create=False")
                
        except Exception as e:
            raise e
    
    def create_index_if_not_exists(self, dimension: int = 1024, metric: str = 'cosine') -> Dict:
        """Створення індексу якщо він не існує (публічний метод)"""
        try:
            existing_indexes = [idx.name for idx in self.pc.list_indexes()]
            
            if self.index_name in existing_indexes:
                return {
                    'success': True,
                    'message': f'Індекс {self.index_name} вже існує',
                    'already_exists': True
                }
            
            # Специфікація для serverless
            spec = ServerlessSpec(
                cloud='aws',
                region='us-east-1'
            )
            
            # Створюємо індекс
            self.pc.create_index(
                name=self.index_name,
                dimension=dimension,
                metric=metric,
                spec=spec
            )
            
            # Оновлюємо підключення
            self.index = self.pc.Index(self.index_name)
            
            # Перевіряємо готовність
            import time
            for i in range(12):  # 60 секунд максимум
                try:
                    stats = self.index.describe_index_stats()
                    return {
                        'success': True,
                        'message': f'Індекс {self.index_name} створено успішно',
                        'dimension': stats.dimension,
                        'created': True
                    }
                except Exception as wait_error:
                    if i < 11:
                        time.sleep(5)
                    else:
                        return {
                            'success': False,
                            'error': f'Timeout waiting for index readiness: {str(wait_error)}'
                        }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def delete_index(self) -> Dict:
        """ОБЕРЕЖНО: Видалення індексу"""
        try:
            self.pc.delete_index(self.index_name)            
            return {
                'success': True,
                'message': f'Індекс {self.index_name} видалено'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def list_all_indexes(self) -> Dict:
        """Список всіх доступних індексів"""
        try:
            indexes = self.pc.list_indexes()
            index_info = []
            
            for idx in indexes:
                try:
                    # Отримуємо детальну інформацію про індекс
                    index_obj = self.pc.Index(idx.name)
                    stats = index_obj.describe_index_stats()
                    
                    info = {
                        'name': idx.name,
                        'dimension': stats.dimension,
                        'total_vectors': stats.total_vector_count,
                        'namespaces': list(stats.namespaces.keys()) if stats.namespaces else [],
                        'status': 'ready'
                    }
                except Exception as index_error:
                    info = {
                        'name': idx.name,
                        'status': 'error',
                        'error': str(index_error)
                    }
                
                index_info.append(info)
            
            return {
                'success': True,
                'indexes': index_info,
                'total_count': len(index_info)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def get_index_info(self) -> Dict:
        """Детальна інформація про поточний індекс"""
        try:
            stats = self.index.describe_index_stats()
            
            # Додаткова інформація про індекс
            index_description = None
            try:
                # Спробуємо отримати опис індексу
                indexes = self.pc.list_indexes()
                for idx in indexes:
                    if idx.name == self.index_name:
                        index_description = {
                            'name': idx.name,
                            'metric': getattr(idx, 'metric', 'unknown'),
                            'spec': str(getattr(idx, 'spec', 'unknown'))
                        }
                        break
            except Exception as desc_error:
                index_description = 'Не вдалося отримати опис індексу'
            
            return {
                'success': True,
                'index_name': self.index_name,
                'dimension': stats.dimension,
                'total_vectors': stats.total_vector_count,
                'namespaces': dict(stats.namespaces) if stats.namespaces else {},
                'index_description': index_description,
                'embedding_model': getattr(self, 'embedding_model', 'unknown'),
                'ready': True
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'index_name': self.index_name,
                'ready': False
            }

    def check_index_status(self) -> Dict:
        """Перевірка стану індексу"""
        try:
            stats = self.index.describe_index_stats()
            
            result = {
                'exists': True,
                'total_vectors': stats.total_vector_count,
                'dimension': stats.dimension,
                'namespaces': dict(stats.namespaces) if stats.namespaces else {},
                'is_empty': stats.total_vector_count == 0
            }
            return result
            
        except Exception as e:
            return {
                'exists': False,
                'error': str(e),
                'is_empty': True
            }
    
    def load_directory(self, directory_path: str, recursive: bool = True) -> Dict:
        """Завантаження всіх файлів з директорії"""
        directory = Path(directory_path)
        
        if not directory.exists() or not directory.is_dir():
            return {'success': False, 'error': f'Директорія не існує: {directory}'}
        
        # Знаходимо всі підтримувані файли
        files_to_process = []

        try:
            if recursive:
                for ext in self.supported_formats.keys():
                    found_files = list(directory.rglob(f"*{ext}"))
                    files_to_process.extend(found_files)
            else:
                for ext in self.supported_formats.keys():
                    found_files = list(directory.glob(f"*{ext}"))
                    files_to_process.extend(found_files)
                        
        except Exception as e:
            return {'success': False, 'error': f'Помилка сканування: {str(e)}'}
        
        if not files_to_process:
            return {'success': False, 'error': 'Не знайдено підтримуваних файлів'}
        
        # Завантажуємо файли
        results = []
        successful = 0
        
        for file_path in enumerate(files_to_process, 1):
            try:
                result = self.load_file(file_path)
                results.append(result)
                
                if result['success']:
                    successful += 1
                    
            except Exception as e:
                error_result = {'success': False, 'error': str(e), 'file': str(file_path)}
                results.append(error_result)
                continue
            
            # Додаємо невелику паузу між файлами для уникнення rate limiting
            import time
            time.sleep(0.1)
        
        return {
            'success': successful > 0,
            'total_files': len(files_to_process),
            'successful': successful,
            'failed': len(files_to_process) - successful,
            'results': results
        }
    
    def load_file(self, file_path: str, source_name: Optional[str] = None) -> Dict:
        """Завантаження одного файлу"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {'success': False, 'error': f'Файл не існує: {file_path}'}
        
        # Визначаємо тип файлу
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_formats:
            return {
                'success': False, 
                'error': f'Непідтримуваний формат: {file_ext}. Підтримувані: {list(self.supported_formats.keys())}'
            }
        
        try:            
            # Витягуємо текст з файлу
            loader_func = self.supported_formats[file_ext]
            text_content = loader_func(file_path)
            
            if not text_content:
                return {'success': False, 'error': 'Файл порожній або не вдалося витягти текст'}
            
            # Розбиваємо на чанки
            chunks = self._split_text(text_content)
            
            if not chunks:
                return {'success': False, 'error': 'Не вдалося створити чанки'}
            
            # Завантажуємо в Pinecone
            source = source_name or file_path.name
            
            result = self._upload_chunks(chunks, source, str(file_path))
            
            if result.get('error'):
                return {
                    'success': False, 
                    'error': result['error'],
                    'partial_upload': result.get('uploaded', 0)
                }
            
            return {
                'success': True,
                'file': str(file_path),
                'source': source,
                'chunks_created': len(chunks),
                'vectors_uploaded': result['uploaded'],
                'text_length': len(text_content)
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _load_pdf(self, file_path: Path) -> str:
        """Завантаження PDF"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Сторінка {page_num + 1} ---\n{page_text}\n"
        except Exception as e:
            logging.error(f"Помилка читання PDF {file_path}: {e}")
        
        return text.strip()
    
    def _load_docx(self, file_path: Path) -> str:
        """Завантаження DOCX"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Також витягуємо текст з таблиць
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"
            
            return text.strip()
            
        except Exception as e:
            return ""
    
    def _load_excel(self, file_path: Path) -> str:
        """Завантаження Excel"""
        try:
            # Читаємо всі листи
            dfs = pd.read_excel(file_path, sheet_name=None)
            text = ""
            
            for sheet_name, df in dfs.items():
                text += f"\n--- Лист: {sheet_name} ---\n"
                text += df.to_string(index=False) + "\n"
            
            return text.strip()
            
        except Exception as e:
            return ""
    
    def _load_csv(self, file_path: Path) -> str:
        """Завантаження CSV"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
        except Exception as e:
            return ""
    
    def _load_epub(self, file_path: Path) -> str:
        """Завантаження EPUB"""
        try:
            book = epub.read_epub(file_path)
            text = ""
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    chapter_text = soup.get_text()
                    if chapter_text.strip():
                        text += chapter_text + "\n\n"
            
            return text.strip()
            
        except Exception as e:
            return ""
    
    def _load_text(self, file_path: Path) -> str:
        """Завантаження текстових файлів"""
        try:
            # Визначаємо кодування
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
            
            # Читаємо файл
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
                
        except Exception as e:
            return ""
    
    def _split_text(self, text: str) -> List[str]:
        """Розбиття тексту на чанки"""
        chunks = []
        
        # Простий алгоритм розбиття по реченнях
        sentences = text.split('.')
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Перевіряємо чи не перевищуємо розмірність чанка
            if len(current_chunk + sentence) < self.chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        # Додаємо останній чанк
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _upload_chunks(self, chunks: List[str], source: str, file_path: str) -> Dict:
        """Завантаження чанків в Pinecone з детальним логуванням"""
        uploaded = 0
        errors = []
        
        try:
            # Перевіряємо підключення перед початком
            stats = self.index.describe_index_stats()
            initial_count = stats.total_vector_count            
            vectors_to_upsert = []
            
            for i, chunk in enumerate(chunks):
                if not chunk.strip():
                    continue
                
                try:
                    # Генеруємо embedding
                    embedding = self._get_embedding(chunk)
                    # Створюємо унікальний ID (ASCII тільки)
                    chunk_id = self._generate_chunk_id(source, i, chunk)
                    # Очищуємо метадані від проблемних символів
                    clean_source = self._clean_text_for_metadata(source)
                    clean_file_path = self._clean_text_for_metadata(file_path)
                    clean_chunk = chunk[:1000]  # Обмежуємо розмір метаданих
                    
                    # Метадані
                    metadata = {
                        'text': clean_chunk,
                        'source': clean_source,
                        'file_path': clean_file_path,
                        'chunk_index': i,
                        'title': f"{clean_source} - частина {i+1}"
                    }
                    
                    vector = {
                        'id': chunk_id,
                        'values': embedding,
                        'metadata': metadata
                    }
                    
                    vectors_to_upsert.append(vector)
                    
                    # Завантажуємо батчами по 100
                    if len(vectors_to_upsert) >= 100:
                        # ДЕТАЛЬНЕ ЛОГУВАННЯ UPSERT
                        try:
                            upsert_response = self.index.upsert(vectors=vectors_to_upsert)
                            uploaded += len(vectors_to_upsert)
                            vectors_to_upsert = []
                            
                            # Перевіряємо чи дійсно додалися вектори
                            import time
                            time.sleep(1)  # Невелика пауза для індексації
                            
                        except Exception as upsert_error:
                            errors.append(f"Upsert error: {str(upsert_error)}")
                            vectors_to_upsert = []  # Очищуємо батч
                            continue
                
                except Exception as chunk_error:
                    error_msg = f"Чанк {i}: {str(chunk_error)}"
                    errors.append(error_msg)
                    continue
            
            # Завантажуємо останні вектори
            if vectors_to_upsert:
                try:
                    uploaded += len(vectors_to_upsert)
                except Exception as final_upsert_error:
                    errors.append(f"Final upsert error: {str(final_upsert_error)}")
            result = {'uploaded': uploaded}
            if errors:
                result['errors'] = errors
                result['error_count'] = len(errors)
            
            return result
            
        except Exception as e:
            return {
                'uploaded': uploaded, 
                'error': str(e),
                'errors': errors
            }
    
    def _clean_text_for_metadata(self, text: str) -> str:
        """Очищення тексту для безпечного збереження в метаданих"""
        import re
        
        # Замінюємо проблемні символи
        clean_text = re.sub(r'[^\w\s\-\.\(\)]+', '_', text)
        
        # Обмежуємо довжину (Pinecone має ліміти на метадані)
        if len(clean_text) > 200:
            clean_text = clean_text[:200] + "..."
        
        return clean_text.strip()
    
    def _detect_embedding_model(self):
        """Автоматичне визначення моделі embedding на основі розмірності індексу"""
        try:
            stats = self.index.describe_index_stats()
            dimension = stats.dimension
            
            if dimension == 1536:
                self.embedding_model = "text-embedding-ada-002"
            elif dimension == 768:
                self.embedding_model = "text-embedding-3-small"
            elif dimension == 3072:
                self.embedding_model = "text-embedding-3-large"
            else:
                # Fallback на локальну модель
                self.embedding_model = "local"
                
        except Exception as e:
            self.embedding_model = "text-embedding-3-small"
    
    def _get_embedding(self, text: str) -> List[float]:
        """Отримання embedding з автоматичним вибором моделі"""
        try:
            # Спочатку перевіряємо розмірність індексу
            stats = self.index.describe_index_stats()
            expected_dim = stats.dimension
            
            if self.embedding_model == "local" or expected_dim == 1024:
                return self._get_local_embedding(text)
            else:
                response = self.openai_client.embeddings.create(
                    input=text,
                    model=self.embedding_model
                )
                embedding = response.data[0].embedding
                
                # Додаткова перевірка розмірності
                if len(embedding) != expected_dim:
                    return self._get_local_embedding(text)
                
                return embedding
                
        except Exception as e:
            return self._get_local_embedding(text)
    
    def _get_local_embedding(self, text: str) -> List[float]:
        """Локальний embedding для нестандартних розмірностей"""
        try:
            
            # Ініціалізуємо модель якщо ще не ініціалізована
            if not hasattr(self, '_local_model'):
                # Модель для української мови з розмірністю 768
                self._local_model = SentenceTransformer('intfloat/multilingual-e5-base')
            
            embedding = self._local_model.encode(text).tolist()
            
            # Перевіряємо розмірність
            stats = self.index.describe_index_stats()
            expected_dim = stats.dimension
            
            if len(embedding) != expected_dim:
                # Обрізаємо або доповнюємо до потрібної розмірності
                if len(embedding) > expected_dim:
                    embedding = embedding[:expected_dim]
                else:
                    embedding.extend([0.0] * (expected_dim - len(embedding)))
            return embedding
            
        except ImportError:
            raise Exception("Локальні embedding недоступні. Встановіть: pip install sentence-transformers")
        except Exception as e:
            raise e
    
    def _generate_chunk_id(self, source: str, chunk_index: int, text: str) -> str:
        """Генерація унікального ID для чанка (тільки ASCII)"""
        
        # Очищуємо source від не-ASCII символів
        clean_source = re.sub(r'[^\x00-\x7F]+', '_', source)  # Замінюємо кирилицю на _
        clean_source = re.sub(r'[^a-zA-Z0-9_\-]', '_', clean_source)  # Тільки букви, цифри, _ та -
        clean_source = clean_source.strip('_')[:50]  # Обмежуємо довжину
        
        # Хеш тексту для унікальності
        text_hash = hashlib.md5(text.encode('utf-8')).hexdigest()[:8]
        
        # Створюємо ASCII ID
        chunk_id = f"{clean_source}_{chunk_index}_{text_hash}"
        
        # Додаткова перевірка - залишаємо тільки ASCII
        chunk_id = ''.join(char for char in chunk_id if ord(char) < 128)
        
        # Якщо ID став порожнім, використовуємо fallback
        if not chunk_id:
            chunk_id = f"doc_{chunk_index}_{text_hash}"
        
        return chunk_id
    
    def clear_index(self) -> Dict:
        """ОБЕРЕЖНО: Очищення всього індексу"""
        try:
            # Отримуємо всі ID
            stats = self.index.describe_index_stats()
            if stats.total_vector_count == 0:
                return {'success': True, 'message': 'Індекс вже порожній'}
            
            # Видаляємо все (це небезпечна операція!)
            self.index.delete(delete_all=True)
            return {
                'success': True, 
                'message': f'Видалено {stats.total_vector_count} векторів'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}