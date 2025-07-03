import ebooklib, fitz, csv, os, re
from ebooklib import epub
from docx import Document
import pandas as pd
from pypdf import PdfReader
from pathlib import Path
from helpers.configuration import Config
config = Config()
def extract_text_from_local_file(file_path):
    ext = Path(file_path).suffix.lower()

    try:
        if ext == ".pdf":
            pdf_reader = PdfReader(file_path)
            text = "\n".join(page.extract_text() for page in pdf_reader.pages if page.extract_text())
            return text

        elif ext in [".docx", ".doc"]:
            doc = Document(file_path)
            return "\n".join(para.text for para in doc.paragraphs)

        elif ext in [".txt", ".md"]:  # <-- додано ".md"
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        elif ext == ".csv":
            df = pd.read_csv(file_path)
            df.dropna(subset=["text"], inplace=True)
            df.drop_duplicates(inplace=True)
            return df.to_csv(index=False)

        elif ext in [".xls", ".xlsx"]:
            df = pd.read_excel(file_path)
            df.dropna(subset=["text"], inplace=True)
            df.drop_duplicates(inplace=True)
            return df.to_csv(index=False)

        else:
            return ""

    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return ""

def normalize_text(text: str) -> str:
    if not text:
        return ""

    # Remove control characters and excessive whitespace
    text = re.sub(r"[\x00-\x1F\x7F]", "", text)  # Remove non-printable characters
    text = re.sub(r"\s+", " ", text)             # Collapse all whitespace to single spaces
    text = re.sub(r"^\s+|\s+$", "", text)        # Strip leading/trailing whitespace

    # Optional: Remove likely metadata lines (e.g., headers/footers from PDFs)
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        # Skip if line is too short or looks like metadata
        if len(line) < 3:
            continue
        if re.search(r"page\s*\d+", line, re.IGNORECASE):
            continue
        if re.match(r"^\W+$", line):  # line is only punctuation or symbols
            continue
        cleaned_lines.append(line)

    cleaned_text = "\n".join(cleaned_lines)

    return cleaned_text

def split_dialog(entry):
    if 'text' in entry and entry['text'] is not None:
        parts = entry['text'].split("### Assistant:\n")
    elif 'text' in entry and entry['text'] is None:
        parts = entry.split("### Assistant:\n")
    elif 'messages' in entry and entry['messages'] is not None:
        parts = entry["messages"]
        parts[0] = f"### User:\n {parts[0]['content'].strip()}"
        parts[1] = parts[1]['content'].strip()
    elif 'prompt' in entry:
        parts = [entry["prompt"].strip(), entry["response"].strip()]
    if len(parts) != 2:
        raise ValueError("Could not split entry into user and assistant parts")
    
    user_part = parts[0].strip()
    assistant_part = parts[1].strip()

    return {
        "prompt": user_part,
        "response": assistant_part
    }

def extract_paragraphs_from_pdf(pdf_path, spltting):
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text() for page in doc])
    return split_to_paragraphs(text, spltting)

def extract_paragraphs_from_docx(docx_path, spltting):
    doc = Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return split_to_paragraphs(text, spltting)

def extract_paragraphs_from_txt(txt_path, spltting):
    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read()
    return split_to_paragraphs(text, spltting)

def extract_paragraphs_from_csv(csv_path, spltting):
    paragraphs = []
    with open(csv_path, newline='', encoding="utf-8") as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            for cell in row:
                if cell.strip():
                    paragraphs.extend(split_to_paragraphs(cell, spltting))
    return paragraphs

def extract_paragraphs_from_excel(excel_path, spltting):
    paragraphs = []
    df = pd.read_excel(excel_path, dtype=str)
    for col in df.columns:
        for cell in df[col].dropna():
            if str(cell).strip():
                paragraphs.extend(split_to_paragraphs(str(cell), spltting))
    return paragraphs

def extract_paragraphs_from_epub(epub_path, spltting):
    book = epub.read_epub(epub_path)
    paragraphs = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            text = item.get_content().decode("utf-8")
            paras = re.findall(r'<p[^>]*>(.*?)</p>', text, re.DOTALL)
            for para in paras:
                clean = re.sub('<[^<]+?>', '', para).strip()
                if clean:
                    paragraphs.append(clean)
    return paragraphs

def split_to_paragraphs(text, splitting=True, min_length=200):
    if not splitting:
        return [text.strip()] if text.strip() else []

    lines = [line.strip() for line in text.split('\n') if line.strip()]
    paragraphs = []
    current = ""

    for line in lines:
        if current:
            current += " " + line
        else:
            current = line
        if len(current) >= min_length:
            paragraphs.append(current.strip())
            current = ""
    if current:
        paragraphs.append(current.strip())

    return paragraphs[:config.MAX_EXAMPLES]

def extract_paragraphs(path, spltting=True):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_paragraphs_from_pdf(path, spltting)
    elif ext == ".docx":
        return extract_paragraphs_from_docx(path, spltting)
    elif ext == ".txt":
        return extract_paragraphs_from_txt(path, spltting)
    elif ext == ".csv":
        return extract_paragraphs_from_csv(path, spltting)
    elif ext in [".xls", ".xlsx"]:
        return extract_paragraphs_from_excel(path, spltting)
    elif ext == ".epub":
        return extract_paragraphs_from_epub(path, spltting)
    else:
        print(f"❌ Непідтримуваний тип файлу: {ext}")
        return []
